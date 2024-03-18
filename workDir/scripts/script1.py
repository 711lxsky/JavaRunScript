import os
from PIL import Image
from docx import Document
from docx.shared import Inches

def find_images(directory):
    """
    在指定目录及其所有子目录中查找图片文件
    """
    img_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
    images = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.lower().endswith(img_extensions):
                images.append(os.path.join(root, file))
    return images

def images_to_word(images, doc_path):
    """
    将图片列表中的所有图片插入到Word文档中
    """
    doc = Document()
    for image_path in images:
        try:
            # 在这里，我们假设所有图片的尺寸都适合放入Word文档中
            # 你可能需要根据图片的实际大小调整插入的大小
            doc.add_picture(image_path, width=Inches(4))
            doc.add_paragraph()  # 为每张图片添加一个空段落以分隔
        except Exception as e:
            print(f"Error adding image {image_path}: {e}")

    doc.save(doc_path)
    print(f"Word document '{doc_path}' has been created with the images.")

# 获取脚本当前的目录
current_directory = os.getcwd()

# 查找当前目录及其子目录下的所有图片
images = find_images(current_directory)

# 将找到的图片插入到Word文档中
word_document_name = "CollectedImages.docx"
images_to_word(images, word_document_name)
